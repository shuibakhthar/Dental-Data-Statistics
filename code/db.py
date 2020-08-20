import sqlite3
import datetime
from datetime import date
from matplotlib import pyplot as plt
import docx
import pandas as pd
import os

class Database:
    def __init__(self,db):
        self.conn=sqlite3.connect(db)
        self.screen = [
            ("Date", "DATE NOT NULL"),
            ("Id_from", "TEXT"),
            ("Id_to", "TEXT"),
            ("TR", "INTEGER"),
            ("Place", "TEXT"),
            ("Address", "TEXT"),
            ("Organiser", "TEXT"),
            ("OPhone", "TEXT"),
            ("Incharge", "TEXT"),
            ("IPhone", "TEXT"),
            ("Endodontics", "INTEGER"),
            ("Periodont", "INTEGER"),
            ("Orthodontics", "INTEGER"),
            ("Surgery", "INTEGER"),
            ("Pedodont", "INTEGER"),
            ("Oralpatho", "INTEGER"),
            ("Publichealth", "INTEGER"),
            ("Oralmeal", "INTEGER"),
            ("Prostho", "INTEGER"),
            ("Total", "INTEGER NOT NULL")
        ]
        self.turnover = [
            ("Id", "TEXT NOT NULL"),
            ("Date", "DATE NOT NULL"),
            ("TR", "INTEGER"),
            ("Incharge", "TEXT"),
            ("IPhone", "TEXT"),
            ("Patient_name", "TEXT"),
            ("Patient_phone", "TEXT"),
            ("H_Education", "INTEGER"),
            ("Endodontics", "INTEGER"),
            ("Periodont", "INTEGER"),
            ("Orthodontics", "INTEGER"),
            ("Surgery", "INTEGER"),
            ("Pedodont", "INTEGER"),
            ("Oralpatho", "INTEGER"),
            ("Publichealth", "INTEGER"),
            ("Oralmeal", "INTEGER"),
            ("Prostho", "INTEGER"),
            ("Total", "INTEGER NOT NULL")
        ]
        self.dept = [
            ("Id", "TEXT NOT NULL"),
            ("Date", "DATE NOT NULL"),
            ("TR", "INTEGER"),
            ("ON_patient", "INTEGER"),
            ("Incharge", "TEXT"),
            ("PhoneNo", "TEXT"),
            ("Patient_name", "TEXT"),
            ("Patient_phone", "TEXT"),
            ("H_Education", "INTEGER"),
            ("Endodontics", "INTEGER"),
            ("Periodont", "INTEGER"),
            ("Orthodontics", "INTEGER"),
            ("Surgery", "INTEGER"),
            ("Pedodont", "INTEGER"),
            ("Oralpatho", "INTEGER"),
            ("Publichealth", "INTEGER"),
            ("Oralmeal", "INTEGER"),
            ("Prostho", "INTEGER"),
            ("Total", "INTEGER")
        ]
        self.satellite = [
            ("Id", "TEXT NOT NULL"),
            ("Date", "DATE NOT NULL"),
            ("TR", "INTEGER"),
            ("Place", "INTEGER"),
            ("Incharge", "TEXT"),
            ("PhoneNo", "TEXT"),
            ("Patient_name", "TEXT"),
            ("Patient_phone", "TEXT"),
            ("Endodontics", "INTEGER"),
            ("Periodont", "INTEGER"),
            ("Orthodontics", "INTEGER"),
            ("Surgery", "INTEGER"),
            ("Pedodont", "INTEGER"),
            ("Oralpatho", "INTEGER"),
            ("Publichealth", "INTEGER"),
            ("Oralmeal", "INTEGER"),
            ("Prostho", "INTEGER"),            
            ("Total", "INTEGER")
        ]
        self.cur=self.conn.cursor()

        # self.drop_table("dept")
        # self.drop_table("turnover")
        # self.drop_table("satellite")
        self.create_table("dept", self.dept)
        self.create_table("screen", self.screen)
        self.create_table("turnover", self.turnover)
        self.create_table("satellite", self.satellite)

    # Utility function to create a table, name with attributes and dtypes in list of tuples as in table
    def create_table(self, name, table, constraint=()):
        query = "CREATE TABLE IF NOT EXISTS %s (%s %s" %(name, table[0][0], table[0][1])
        for row in table[1:]:
            query += ", %s %s" %(row[0], row[1])
        query += ')'
        # if constraint == ():
        #     query += ");"
        # elif type(constraint) == str:
        #     query += ", PRIMARY KEY (%s))" %constraint
        # else:
        #     query += ", CONSTRAINT Con_prime PRIMARY KEY (" 
        #     query += ("%s, " * (len(constraint)-1) +"%s));") %constraint
        self.cur.execute(query)
        self.conn.commit()
    
    def drop_table(self, name):
        self.cur.execute("DROP TABLE %s" %name)
        self.conn.commit()

    def get_count(self, name):
        self.cur.execute("SELECT COUNT(*) FROM %s" %(name, ))
        count = self.cur.fetchall()
        return count[0][0]


    def fetch(self, name):
        self.cur.execute("SELECT * FROM %s" %(name, ))
        rows = self.cur.fetchall()
        return rows
    
    # Returns the date, total and treated status in between given dates
    def fetch_by_date(self, name, from_date, to_date, TR = 0):
        query = ""
        if TR == 0:
            query = "SELECT Date, sum(total) FROM %s WHERE Date > '%s' AND Date < '%s' GROUP BY Date ORDER BY Date" %(name, from_date, to_date)
        else:
            query = "SELECT Date, sum(total) FROM %s WHERE Date > '%s' AND Date < '%s' AND TR = %d GROUP BY Date ORDER BY Date" %(name, from_date, to_date, TR)
        # print(query)
        self.cur.execute(query)
        rows = self.cur.fetchall()
        # print(rows)
        cols = list(zip(*rows))
        # print(cols)
        return cols

    def fetch_for_word(self,date):
        self.cur.execute("SELECT Date, Organiser, Ophone, Incharge, Iphone, Total FROM screen WHERE Date ='%s'"%(date))
        wordvalues=self.cur.fetchall()
        wordlist=[]
        for i in wordvalues:
            for j in i:
                wordlist.append(j)
        # print(wordlist)
        return wordlist


    # Utility function to insert a row in the table for the list of values
    def insert(self, name, values):
        query = ""
        if name == 'screen':
            query += "insert into screen values(" + "?," *(len(self.screen)-1) + "?)"
            values = tuple(values)
        elif name == 'turnover':
            query += "insert into turnover values(" + "?," *(len(self.turnover)-1) + "?)"
            values = tuple(values)
        elif name == 'dept':
            query += "insert into dept values(" + "?," *(len(self.dept)-1) + "?)"
            values = tuple(values)
        elif name == 'satellite':
            query += "insert into satellite values(" + "?," *(len(self.satellite)-1) + "?)"
            values = tuple(values)
        self.cur.execute(query, values)
        self.conn.commit()

    # Deleting a record with the given date
    def remove(self, name, del_id1, del_id2 = ''):
        if name == "screen":
            query = "DELETE FROM %s WHERE Date='%s'" %(name, del_id1)
        else:
            query = "DELETE FROM %s WHERE Id='%s' AND Date='%s'" %(name, del_id1, del_id2)
        self.cur.execute(query)
        self.conn.commit()

    # Updates the value of given table-name 
    def update(self, name, new_values):
        # print(new_values, type(new_values))
        if name == "screen":
            n = new_values[0]
            new_values = new_values[1:]
            new_values.append(n)
        else:
            n,m = new_values[:2]
            new_values = new_values[2:]
            new_values.append(n)
            new_values.append(m)
        # print(new_values)
        if name == 'dept':
            attrib = [attr[0] for attr in self.dept[2:len(self.dept)]]            
        elif name == 'screen':
            attrib = [attr[0] for attr in self.screen[1:len(self.screen)]]
        elif name == 'turnover':
            attrib = [attr[0] for attr in self.turnover[2:len(self.turnover)]]
        elif name == "satellite":
            attrib = [attr[0] for attr in self.satellite[2:len(self.satellite)]]
        # print(attrib)
        query = "UPDATE %s SET "%(name)
        query += ("%s = ?, " *(len(attrib)-1) + "%s = ?") %tuple(attrib)
        if name == 'screen':
            query += " WHERE Date = ?"
        else:
            query +=  " WHERE Id = ? AND Date = ?" 
        # print(query)
        self.cur.execute(query, new_values)
        self.conn.commit()
    
    # Gets the count of all values in  column names given in sums in a particular date from table name given
    def get_total(self, name, sums, date, TR):
        counts = []
        query = "SELECT "
        query += ("SUM(%s), " * (len(sums)-1) + " SUM(%s)") %tuple(i[0] for i in sums)
        query += " FROM %s WHERE Date = '%s' AND TR = %d" %(name, date, TR)
        # print(query)
        self.cur.execute(query)
        rows = self.cur.fetchall()
        # print(rows)
        self.conn.commit()
        return rows[0]

    def get_total1(self, name, sums, date, TR,place):
        counts = []
        query = "SELECT "
        query += ("SUM(%s), " * (len(sums)-1) + " SUM(%s)") %tuple(i[0] for i in sums)
        query += " FROM %s WHERE (Date = '%s' AND TR = %d AND place = %d)" %(name, date, TR, place)
        # print(query)
        self.cur.execute(query)
        rows = self.cur.fetchall()
        # print(rows)
        self.conn.commit()
        return rows[0]

    def dept_func(self, date):
        self.cur.execute("""select count(ID) from dept where Date = :1 """, {"1": date})
        tot = self.cur.fetchall()
        self.cur.execute("""select count(ID) from dept where Date = :1 and H_Education = 1""", {"1": date})
        h = self.cur.fetchall()
        self.cur.execute("""select count(ID) from dept where Date = :1 and ON_patient = 1""", {"1": date})
        old = self.cur.fetchall()
        self.cur.execute("""select count(ID) from dept where Date = :1 and ON_patient = 2""", {"1": date})
        new = self.cur.fetchall()
        self.cur.execute("""select count(ID) from dept where Date = :1 and TR = 1""", {"1": date})
        s = self.cur.fetchall()
        self.cur.execute("""select count(ID) from dept where Date = :1 and TR = 2""", {"1": date})
        r = self.cur.fetchall()
        self.cur.execute("""select count(ID) from dept where Date = :1 and TR = 3""", {"1": date})
        t = self.cur.fetchall()
        # print()
        return tot[0],h[0],old[0], new[0], s[0], r[0], t[0]

    def sat_func(self,date,centre):
        # print("HI")
        self.cur.execute("""select count(ID) from satellite where Date = :1 and Place = :2 """, {"1": date, "2": centre})
        tot = self.cur.fetchall()

        self.cur.execute("""select count(ID) from satellite where (Date = :1 and TR =1 and Place = :2)""", {"1": date, "2": centre})
        s = self.cur.fetchall()

        self.cur.execute("""select count(ID) from satellite where (Date = :1 and TR =2 and Place = :2)""", {"1": date, "2": centre})
        r = self.cur.fetchall()

        self.cur.execute("""select count(ID) from satellite where (Date = :1 and TR =3 and Place = :2)""", {"1": date, "2": centre})
        t = self.cur.fetchall()

        return tot[0], s[0], r[0], t[0]

    def turn_Func(self,date):
        # print("HI"
        self.cur.execute("""select count(ID) from turnover where Date = :1 """, {"1": date})
        tot = self.cur.fetchall()

        self.cur.execute("""select count(ID) from turnover where Date = :1 and H_Education =1""", {"1": date})
        h = self.cur.fetchall()

        self.cur.execute("""select count(ID) from turnover where Date = :1 and TR =1""", {"1": date})
        s = self.cur.fetchall()

        self.cur.execute("""select count(ID) from turnover where Date = :1 and TR =2""", {"1": date})
        r = self.cur.fetchall()

        self.cur.execute("""select count(ID) from turnover where Date = :1 and TR =3""", {"1": date})
        t = self.cur.fetchall()



        return tot[0],h[0],s[0],r[0],t[0]

    def __del__(self):
        self.conn.close()

    def createbar(self):
        xvalues,yvalues = self.fetch_by_date("dept", "1999-01-01", "2005-01-01")
        plt.bar(xvalues,yvalues)
        plt.xlabel("Date")
        plt.ylabel("Patients")
        plt.show()

    def createline(self):
        xvalues,yvalues = self.fetch_by_date("dept", "1999-01-01", "2005-01-01")
        plt.plot(xvalues,yvalues)
        plt.xlabel("Date")
        plt.ylabel("Patients")
        plt.grid()
        plt.show()

    # To get the total catagorised by month
    def get_by_date(self, name, from_date, to_date, group_by='m'):
        monthDict = {'01': 'Jan', '02': 'Feb', '03': 'Mar', '04': 'Apr', '05': 'May', '06': 'Jun', '07': 'Jul',
                     '08': 'Aug', '09': 'Sep', '10': 'Oct', '11': 'Nov', '12': 'Dec'}
        if group_by == 'm' or group_by == 'M':
            query = "SELECT strftime('%m-%Y', Date) as mon_date, sum(Total) FROM"
        elif group_by == 'Y' or group_by == 'y':
            query = "SELECT strftime('%Y', Date) as mon_date, sum(Total) FROM"
        query += " %s WHERE Date BETWEEN '%s' and '%s' GROUP BY mon_date ORDER BY mon_date" % (
        name, from_date, to_date)
        # print(query)
        self.cur.execute(query)
        self.conn.commit()
        rows = self.cur.fetchall()
        cols = list(zip(*rows))
        # cols = [[monthDict[i] for i in cols[0]], cols[1]]
        return tuple(cols)

    # To create a word doc
    def createword(self,date):
        doc=docx.Document()
        #obj1=doc.paragraphs[0].text;
        #obj2=doc.paragraphs[1].text;
        #doc.paragraphs[0].text=obj1+"12/12/12";
        #doc.paragraphs[1].text=obj2+"12:00";
        #doc.add_paragraph(doc.paragraphs)
        values=self.fetch_for_word(date)
        # print(values)
        #values=list(values)
        #print(values)
        dateobj=doc.add_paragraph("Date:  "+values[0])
        #timeobj=doc.add_paragraph("Time: "+values[1])
        incharge=doc.add_paragraph("Organizer: "+values[1])
        internobj=doc.add_paragraph("Organizer Phone "+values[2])
        organizer=doc.add_paragraph("Incharge: "+values[3])
        patient=doc.add_paragraph("Incharge Phone  "+values[4])
        #systempath=os.getcwd()
        #newpath=systempath+"/"+"images/"+date
        #import os
        systempath=os.getcwd()
        newpath=systempath+"/"+"images/"+date
        # print(newpath)
        directory = os.fsencode(newpath)
        desktop=os.path.expanduser("~/Desktop")
        # print(desktop)
        for file in os.listdir(directory):
            filename = os.fsdecode(file)
            if filename.endswith(".jpg") or filename.endswith(".png"):
                try:
                 doc.add_picture(newpath+"\{}".format(filename))
                except FileNotFoundError: 
                 continue
            else:
                 continue
        fn="%s"%(date)
        doc.save(fn+".docx")
        os.rename(fn+".docx",desktop+"/{}".format(date)+".docx")
    
    # To create an excel file
    def createexcel(self, tablename, fromdate, todate, TR=0):
        conn = sqlite3.connect("dental.db")
        if TR == 0:
            df = pd.read_sql_query(
                """ select * from %s  WHERE Date > '%s' AND Date < '%s' GROUP BY Date ORDER BY Date""" % (
                tablename, fromdate, todate), conn)
        else:
            df = pd.read_sql_query(
                """SELECT * FROM %s WHERE Date > '%s' AND Date < '%s' AND TR = %d GROUP BY Date ORDER BY Date""" % (
                tablename, fromdate, todate, TR), conn)
        desktop = os.path.expanduser("~/Desktop")
        fn="%s"%(tablename)
        df.to_excel("{}".format(tablename) + ".xlsx")
        os.rename(fn + ".xlsx", desktop + "/{}".format(tablename) + ".xlsx")




db = Database("dental.db")
# db.get_total("dept",db.dept[8:], '2020-03-10',3)
# # db.insert("dept", )
# db.get_total("dept", db.dept[7:], "1999-10-11", 1)

#db.createbar()
#db.createline()
# row = db.fetch("dept")
# print(row)
# print()
# total = db.fetch_by_date("dept", "1999-01-01", "2005-01-01", 2)
# print(total)

# if __name__ == "__main__":
#     db = Database("dental.db")
#     db.insert("patient",
#         (   "10",
#             "SRM axis",
#             "Ramapuram",
#             "7799662288",
#             19,
#             4987,
#             4987,
#             4987,
#             4987,
#             4987,
#             4987,
#             4987,
#             124
#         )
#     )
#     print(db.fetch("patient"))
#     db.update("patient",
#         (   "10",
#             "SRM dental",
#             "Ramapuram",
#             "7799662288",
#             19,
#             4987,
#             4987,
#             4987,
#             4987,
#             4987,
#             4987,
#             4987,
#             124
#         )
#     )
#     print(db.fetch("patient"))
